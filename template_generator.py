from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import tempfile
import os
from datetime import datetime

class AzureTemplateGenerator:
    """Generates Azure Synapse Pipeline documentation using Word templates"""
    
    def __init__(self):
        self.doc = None
    
    def generate_document(self, synthesis_data, uploaded_screenshots=None):
        """Generate complete Azure Synapse Pipeline documentation"""
        try:
            # Create new document
            self.doc = Document()
            
            # Store screenshots for later use
            self.screenshots = uploaded_screenshots or []
            
            # Add title
            self._add_title()
            
            # Add sections based on the template
            self._add_general_information(synthesis_data)
            self._add_sources_and_sinks(synthesis_data)
            self._add_data_flow_logic(synthesis_data)
            self._add_dependencies(synthesis_data)
            self._add_performance_maintenance(synthesis_data)
            self._add_change_log()
            
            # Save document
            doc_path = self._save_document()
            return doc_path
            
        except Exception as e:
            print(f"Error generating document: {str(e)}")
            return None
    
    def _add_title(self):
        """Add document title"""
        title = self.doc.add_heading('Azure Synapse Pipeline – Hospital Inpatient Dashboard - Documentation', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()
    
    def _add_general_information(self, synthesis_data):
        """Add General Information section"""
        self.doc.add_heading('1. General Information', level=1)
        
        # Data Flow Name subsection
        self.doc.add_heading('Data Flow Name:', level=2)
        
        explanation = self.doc.add_paragraph()
        explanation.add_run(
            "Provide the official name of the data flow as it appears in Azure Synapse Analytics. "
            "This should be specific enough to distinguish it from other data flows in your environment. "
            "Use your organization's naming conventions to ensure consistency."
        )
        
        # Extract data flow name from synthesis
        data_flow_name = self._extract_data_flow_name(synthesis_data)
        name_para = self.doc.add_paragraph()
        name_para.add_run(f"Data Flow Name: {data_flow_name}").bold = True
        
        self.doc.add_paragraph()
        
        # Data Flow Screenshot subsection
        self.doc.add_heading('Data Flow Screenshot:', level=2)
        
        # Add screenshots if provided
        if hasattr(self, 'screenshots') and self.screenshots:
            for i, screenshot in enumerate(self.screenshots, 1):
                try:
                    # Add screenshot caption
                    caption_para = self.doc.add_paragraph()
                    caption_para.add_run(f"Screenshot {i}: {screenshot.name}").bold = True
                    
                    # Save screenshot temporarily and add to document
                    import tempfile
                    with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{screenshot.name}") as tmp_file:
                        tmp_file.write(screenshot.getvalue())
                        tmp_file_path = tmp_file.name
                    
                    # Add image to document (resize to fit page)
                    try:
                        self.doc.add_picture(tmp_file_path, width=Inches(6.0))
                    except Exception as img_error:
                        print(f"Error adding image {screenshot.name}: {img_error}")
                        self.doc.add_paragraph(f"[Screenshot: {screenshot.name} - Could not embed image]")
                    
                    # Clean up temporary file
                    os.unlink(tmp_file_path)
                    
                    self.doc.add_paragraph()  # Add spacing
                    
                except Exception as e:
                    print(f"Error processing screenshot {screenshot.name}: {str(e)}")
                    self.doc.add_paragraph(f"[Screenshot: {screenshot.name} - Error processing image]")
        else:
            self.doc.add_paragraph("[Screenshot placeholder - Please insert actual data flow screenshot here]")
        
        self.doc.add_paragraph()
    
    def _add_sources_and_sinks(self, synthesis_data):
        """Add Sources and Sinks section"""
        self.doc.add_heading('2. Source(s) and Sink(s)', level=1)
        
        # Create table
        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['Source/Sink Name', 'Type', 'Location', 'Format', 'Description/Notes']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # Extract sources and sinks from synthesis
        sources_sinks = self._extract_sources_and_sinks(synthesis_data)
        
        # Add rows for sources and sinks
        for item in sources_sinks:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get('name', 'N/A')
            row_cells[1].text = item.get('type', 'N/A')
            row_cells[2].text = item.get('location', 'N/A')
            row_cells[3].text = item.get('format', 'N/A')
            row_cells[4].text = item.get('description', 'N/A')
        
        self.doc.add_paragraph()
    
    def _add_data_flow_logic(self, synthesis_data):
        """Add Data Flow Logic & Business Rules section"""
        self.doc.add_heading('3. Data Flow Logic & Business Rules', level=1)
        
        # Step-by-Step Breakdown
        self.doc.add_heading('Step-by-Step Breakdown:', level=2)
        
        explanation = self.doc.add_paragraph()
        explanation.add_run(
            "Document complex transformation steps in the data flow if its used, "
            "no need to capture simple transformation like joins / unions etc.,"
        )
        
        # Extract transformations from synthesis
        transformations = self._extract_transformations(synthesis_data)
        
        for transformation in transformations:
            para = self.doc.add_paragraph()
            para.add_run(f"{transformation['name']}: ").bold = True
            para.add_run(transformation['description'])
        
        self.doc.add_paragraph()
        
        # Key Transformations/Business Rules
        self.doc.add_heading('Key Transformations/Business Rules:', level=2)
        
        business_rules = self._extract_business_rules(synthesis_data)
        
        for i, rule in enumerate(business_rules, 1):
            para = self.doc.add_paragraph()
            para.add_run(f"Rule {i}: ").bold = True
            para.add_run(rule)
        
        explanation2 = self.doc.add_paragraph()
        explanation2.add_run(
            "Expand upon each rule, describing its business impact, rationale, and technical implementation."
        )
        
        self.doc.add_paragraph()
    
    def _add_dependencies(self, synthesis_data):
        """Add Dependencies & Related Assets section"""
        self.doc.add_heading('4. Dependencies & Related Assets', level=1)
        
        # Upstream Dependencies
        self.doc.add_heading('Upstream Dependencies:', level=2)
        
        explanation = self.doc.add_paragraph()
        explanation.add_run(
            "Identify the ETL pipelines, data ingestion processes, or external systems supplying "
            "source data to the data flow. Document their names, owners, schedules, and any critical "
            "SLAs (Service Level Agreements) or data freshness requirements."
        )
        
        upstream_deps = self._extract_upstream_dependencies(synthesis_data)
        if upstream_deps:
            for dep in upstream_deps:
                self.doc.add_paragraph(f"• {dep}", style='List Bullet')
        else:
            self.doc.add_paragraph("No upstream dependencies identified from the provided documentation.")
        
        self.doc.add_paragraph()
        
        # Downstream Dependencies
        self.doc.add_heading('Downstream Dependencies:', level=2)
        
        explanation2 = self.doc.add_paragraph()
        explanation2.add_run(
            "List reports, dashboards (e.g., Power BI), machine learning models, or other business "
            "processes that consume the output. Clarify how changes to this data flow might impact "
            "these dependencies, and note stakeholders for coordination."
        )
        
        downstream_deps = self._extract_downstream_dependencies(synthesis_data)
        if downstream_deps:
            for dep in downstream_deps:
                self.doc.add_paragraph(f"• {dep}", style='List Bullet')
        else:
            self.doc.add_paragraph("No downstream dependencies identified from the provided documentation.")
        
        self.doc.add_paragraph()
    
    def _add_performance_maintenance(self, synthesis_data):
        """Add Performance & Maintenance section"""
        self.doc.add_heading('6. Performance & Maintenance', level=1)
        
        # Integration Runtime
        self.doc.add_heading('Integration Runtime:', level=2)
        
        explanation = self.doc.add_paragraph()
        explanation.add_run(
            "Specify which Azure Integration Runtime is used (auto or self-hosted), compute size, "
            "and any configuration details that affect execution or scalability."
        )
        
        integration_runtime = self._extract_integration_runtime(synthesis_data)
        self.doc.add_paragraph(integration_runtime)
        
        self.doc.add_paragraph()
        
        # Known Issues/Considerations
        self.doc.add_heading('Known Issues/Considerations:', level=2)
        
        explanation2 = self.doc.add_paragraph()
        explanation2.add_run(
            "Document any limitations, data quality concerns, error patterns, or historical "
            "challenges encountered in this data flow. Include troubleshooting tips if available."
        )
        
        known_issues = self._extract_known_issues(synthesis_data)
        if known_issues:
            for issue in known_issues:
                self.doc.add_paragraph(f"• {issue}", style='List Bullet')
        else:
            self.doc.add_paragraph("No specific issues identified from the provided documentation.")
        
        self.doc.add_paragraph()
        
        # Performance Tuning Notes
        self.doc.add_heading('Performance Tuning Notes:', level=2)
        
        explanation3 = self.doc.add_paragraph()
        explanation3.add_run(
            "Record optimizations made to improve throughput, latency, or reliability. "
            'For example: "Increased parallelism by scaling out to 16 cores," or '
            '"Partitioned source files by month to reduce load time."'
        )
        
        perf_notes = self._extract_performance_notes(synthesis_data)
        if perf_notes:
            for note in perf_notes:
                self.doc.add_paragraph(f"• {note}", style='List Bullet')
        else:
            self.doc.add_paragraph("No specific performance tuning notes identified from the provided documentation.")
        
        self.doc.add_paragraph()
    
    def _add_change_log(self):
        """Add Change Log section"""
        self.doc.add_heading('7. Change Log', level=1)
        
        # Create change log table
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['Version', 'Date', 'Author', 'Changes Made']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # Add initial version
        row_cells = table.add_row().cells
        row_cells[0].text = '1.0'
        row_cells[1].text = datetime.now().strftime('%Y-%m-%d')
        row_cells[2].text = 'DocuFillGenie AI'
        row_cells[3].text = 'Initial version generated from uploaded documentation.'
        
        self.doc.add_paragraph()
    
    def _extract_data_flow_name(self, synthesis_data):
        """Extract data flow name from synthesis"""
        # First check for structured data from LLM analysis
        structured_data = synthesis_data.get('structured_data', {})
        pipeline_names = structured_data.get('pipeline_names', [])
        
        if pipeline_names:
            # Use the first pipeline name found
            return pipeline_names[0]
        
        # Fall back to synthesis text analysis
        synthesis = synthesis_data.get('synthesis', '')
        
        # Look for pipeline names in the synthesis text
        import re
        pipeline_patterns = [
            r'Main Data Flow:\s*([^\n]+)',
            r'Pipeline Names Identified:\s*([^\n]+)',
            r'([A-Za-z0-9\s]+(?:Dashboard|Pipeline|Data Flow))',
        ]
        
        for pattern in pipeline_patterns:
            matches = re.findall(pattern, synthesis, re.IGNORECASE)
            if matches:
                return matches[0].strip()
        
        # Check source file names for pipeline hints
        source_files = synthesis_data.get('source_files', [])
        for filename in source_files:
            if 'dashboard' in filename.lower():
                # Extract meaningful name from filename
                name_part = filename.replace('.zip', '').replace('.docx', '').replace('_', ' ')
                return f"{name_part} Data Flow"
        
        return "Azure Synapse Data Flow"
    
    def _extract_sources_and_sinks(self, synthesis_data):
        """Extract sources and sinks information"""
        sources_sinks = []
        
        # Get structured data from LLM analysis
        structured_data = synthesis_data.get('structured_data', {})
        sources = structured_data.get('sources', [])
        sinks = structured_data.get('sinks', [])
        
        # Add identified sources
        for i, source in enumerate(sources[:3], 1):  # Limit to 3 sources
            source_type = 'Database'
            source_format = 'SQL Server'
            
            # Determine type and format based on source name
            source_lower = source.lower()
            if any(ext in source_lower for ext in ['.csv', '.excel', '.xlsx']):
                source_type = 'File'
                source_format = 'Excel/CSV'
            elif any(db in source_lower for db in ['sql', 'database', 'db']):
                source_type = 'Database'
                source_format = 'SQL Server'
            elif any(api in source_lower for api in ['api', 'service', 'endpoint']):
                source_type = 'API'
                source_format = 'REST API'
            
            sources_sinks.append({
                'name': source,
                'type': source_type,
                'location': 'As specified in source documentation',
                'format': source_format,
                'description': f'Data source identified from uploaded documentation'
            })
        
        # Add identified sinks
        for i, sink in enumerate(sinks[:3], 1):  # Limit to 3 sinks
            sink_type = 'Data Warehouse'
            sink_format = 'Parquet'
            
            # Determine type and format based on sink name
            sink_lower = sink.lower()
            if 'synapse' in sink_lower or 'warehouse' in sink_lower:
                sink_type = 'Data Warehouse'
                sink_format = 'Parquet'
            elif 'power bi' in sink_lower or 'powerbi' in sink_lower:
                sink_type = 'BI Platform'
                sink_format = 'Dataset'
            elif 'dashboard' in sink_lower:
                sink_type = 'Dashboard'
                sink_format = 'Aggregated Data'
            
            sources_sinks.append({
                'name': sink,
                'type': sink_type,
                'location': 'As specified in destination documentation',
                'format': sink_format,
                'description': f'Data destination identified from uploaded documentation'
            })
        
        # If no sources/sinks found, add at least one placeholder with better naming
        if not sources_sinks:
            pipeline_name = self._extract_data_flow_name(synthesis_data)
            
            sources_sinks.extend([
                {
                    'name': f'{pipeline_name} Source',
                    'type': 'Database',
                    'location': 'To be specified from documentation',
                    'format': 'SQL Server',
                    'description': 'Primary data source - requires manual specification from detailed documentation'
                },
                {
                    'name': f'{pipeline_name} Destination',
                    'type': 'Data Warehouse',
                    'location': 'Azure Synapse',
                    'format': 'Parquet',
                    'description': 'Primary destination - requires manual specification from detailed documentation'
                }
            ])
        
        return sources_sinks
    
    def _extract_transformations(self, synthesis_data):
        """Extract transformation steps"""
        structured_data = synthesis_data.get('structured_data', {})
        transformations_list = structured_data.get('transformations', [])
        
        transformations = []
        
        # Use actual transformations found in documents
        for i, transform in enumerate(transformations_list[:5], 1):  # Limit to 5
            transformations.append({
                'name': f'Transform_{i}_{transform.replace(" ", "_")}',
                'description': f'Process and transform data through {transform} as identified in the source documentation.'
            })
        
        # If no specific transformations found, add generic ones based on pipeline name
        if not transformations:
            pipeline_name = self._extract_data_flow_name(synthesis_data)
            
            transformations.extend([
                {
                    'name': 'Filter_ValidRecords',
                    'description': f'Filter and validate records for {pipeline_name} processing.'
                },
                {
                    'name': 'DerivedColumn_Calculations',
                    'description': f'Compute derived metrics and calculated fields for {pipeline_name} as identified in the source documentation.'
                }
            ])
        
        return transformations
    
    def _extract_business_rules(self, synthesis_data):
        """Extract business rules"""
        structured_data = synthesis_data.get('structured_data', {})
        business_rules_list = structured_data.get('business_rules', [])
        pipeline_name = self._extract_data_flow_name(synthesis_data)
        
        rules = []
        
        # Use actual business rules found in documents
        for rule in business_rules_list[:5]:  # Limit to 5 rules
            rules.append(rule)
        
        # Add some standard rules if none found
        if not rules:
            rules.extend([
                f"Apply data quality validation for {pipeline_name} to ensure all required fields are populated.",
                f"Transform date formats for {pipeline_name} to comply with organizational standards.",
                f"Implement lookup transformations for {pipeline_name} reference data as specified in the documentation."
            ])
        
        return rules
    
    def _extract_upstream_dependencies(self, synthesis_data):
        """Extract upstream dependencies"""
        source_files = synthesis_data.get('source_files', [])
        deps = []
        
        for file in source_files:
            deps.append(f"Source documentation: {file}")
        
        deps.extend([
            "Data ingestion processes feeding the identified sources",
            "External systems providing reference data"
        ])
        
        return deps
    
    def _extract_downstream_dependencies(self, synthesis_data):
        """Extract downstream dependencies"""
        return [
            "Hospital Inpatient Dashboard (Power BI)",
            "Reporting systems consuming processed data",
            "Data mart serving analytical queries"
        ]
    
    def _extract_integration_runtime(self, synthesis_data):
        """Extract integration runtime information"""
        return "Azure Integration Runtime (Auto-resolve) - Default configuration with auto-scaling enabled."
    
    def _extract_known_issues(self, synthesis_data):
        """Extract known issues"""
        return [
            "Monitor for data latency during peak processing hours",
            "Validate data quality metrics as defined in the source documentation"
        ]
    
    def _extract_performance_notes(self, synthesis_data):
        """Extract performance tuning notes"""
        return [
            "Optimize for parallel processing based on data volume patterns",
            "Consider partitioning strategies for large datasets"
        ]
    
    def _save_document(self):
        """Save the document to a temporary file"""
        try:
            # Create temporary file
            temp_dir = tempfile.gettempdir()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Azure_Synapse_Pipeline_Documentation_{timestamp}.docx"
            file_path = os.path.join(temp_dir, filename)
            
            # Save document
            self.doc.save(file_path)
            return file_path
            
        except Exception as e:
            print(f"Error saving document: {str(e)}")
            return None
