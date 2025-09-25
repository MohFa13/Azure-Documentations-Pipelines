import os
import zipfile
import tempfile
from pathlib import Path
import pandas as pd
from docx import Document
import PyPDF2
# import magic  # Optional dependency, will use file extension fallback
import chardet
import json
import re
from difflib import get_close_matches

class DocumentProcessor:
    """Handles extraction and processing of various document formats"""
    
    def __init__(self):
    """Initialize the document processor"""
    self.supported_formats = ['.pdf', '.docx', '.txt', '.json', '.zip']
    
    # Database mapping from your image
    self.db_mapping = {
        'agilist_datalab_test': 'Oracle',
        'AGILIST': 'Oracle', 
        'AGILIST_DATALAB': 'Oracle',
        'AGILIST_DELV': 'Oracle',
        'AGLPRD': 'Oracle',
        'AGLPRD_IM_ATM': 'Oracle',
        'AGLPRDDELV': 'Oracle',
        'AGLPRDIMSTAGE': 'Oracle',
        'AzureDataLakeStorage1': 'Azure Data Lake Storage Gen2',
        'BOXI': 'Oracle',
        'Dataprd': 'SQL Server',
        'prdpi': 'Oracle',
        'ResusApp': 'SQL Server',
        'SHDSQLPRD11': 'SQL Server',
        # Additional common patterns from your JSON example
        'AGLPRD24': 'Oracle',
        'PipelineParquet250': 'Azure Data Lake Storage Gen2',
        'synw-prod-qc-01-WorkspaceDefaultStorage': 'Azure Data Lake Storage Gen2'
    }

def process_zip_file(self, uploaded_file):
    """Process zip file containing folders with JSON files, focusing on pipeline folder"""
    try:
        # Read the zip file
        zip_bytes = BytesIO(uploaded_file.getvalue())
        extracted_data = []
        
        with zipfile.ZipFile(zip_bytes, 'r') as zip_ref:
            # Get all files in the zip
            all_files = zip_ref.namelist()
            
            # Filter for JSON files specifically in pipeline folder
            pipeline_json_files = [
                f for f in all_files 
                if 'pipeline' in f.lower() and f.endswith('.json') and not f.startswith('__MACOSX')
            ]
            
            if not pipeline_json_files:
                # Fallback: look for any JSON files
                json_files = [f for f in all_files if f.endswith('.json') and not f.startswith('__MACOSX')]
                if json_files:
                    st.warning(f"⚠️ No files found in 'pipeline' folder. Processing {len(json_files)} JSON files from other locations.")
                    pipeline_json_files = json_files
                else:
                    return [{'error': 'No JSON files found in the ZIP file'}]
            
            st.info(f"📄 Found {len(pipeline_json_files)} JSON files to process")
            
            for json_file in pipeline_json_files:
                try:
                    # Read JSON content
                    with zip_ref.open(json_file) as file:
                        json_content = json.loads(file.read().decode('utf-8'))
                    
                    # Extract folder and file info
                    folder_path = os.path.dirname(json_file)
                    file_name = os.path.basename(json_file)
                    
                    # Process the pipeline data
                    pipeline_data = self.extract_pipeline_data(json_content)
                    
                    # Add metadata to each activity
                    for activity in pipeline_data:
                        if 'error' not in activity:
                            activity['source_folder'] = folder_path
                            activity['source_file'] = file_name
                            activity['pipeline_name'] = json_content.get('name', file_name.replace('.json', ''))
                    
                    extracted_data.extend(pipeline_data)
                    
                except Exception as e:
                    extracted_data.append({
                        'error': f'Error processing {json_file}: {str(e)}',
                        'source_folder': os.path.dirname(json_file),
                        'source_file': os.path.basename(json_file)
                    })
        
        return extracted_data
        
    except Exception as e:
        return [{'error': f'Error processing zip file: {str(e)}'}]

def _process_docx(self, file_path: str) -> str:
    """Process DOCX file - implement your DOCX processing logic here"""
    try:
        # You'll need to install and import python-docx
        # from docx import Document
        # doc = Document(file_path)
        # text = ""
        # for paragraph in doc.paragraphs:
        #     text += paragraph.text + "\n"
        # return text
        
        return "DOCX processing not implemented yet - add python-docx library"
    except Exception as e:
        return f"Error processing DOCX: {str(e)}"

def _process_txt(self, uploaded_file) -> str:
    """Process TXT file"""
    try:
        return uploaded_file.getvalue().decode("utf-8")
    except Exception as e:
        return f"Error processing TXT: {str(e)}"

def extract_pipeline_data(self, json_content):
    """Extract pipeline data from Azure Data Factory JSON"""
    try:
        # Parse JSON if it's a string
        if isinstance(json_content, str):
            pipeline_data = json.loads(json_content)
        else:
            pipeline_data = json_content
        
        activities = pipeline_data.get('properties', {}).get('activities', [])
        extracted_data = []
        
        for activity in activities:
            activity_info = {
                'activity_name': activity.get('name', ''),
                'activity_type': activity.get('type', ''),
                'source_name': '',
                'source_type': '',
                'sink_name': '',
                'sink_type': '',
                'depends_on': [],
                'timeout': '',
                'retry': 0
            }
            
            # Extract dependency information
            depends_on = activity.get('dependsOn', [])
            if depends_on:
                activity_info['depends_on'] = [dep.get('activity', '') for dep in depends_on]
            
            # Extract policy information
            policy = activity.get('policy', {})
            if policy:
                activity_info['timeout'] = policy.get('timeout', '')
                activity_info['retry'] = policy.get('retry', 0)
            
            # Extract source information
            if 'inputs' in activity and activity['inputs']:
                source_ref = activity['inputs'][0].get('referenceName', '')
                activity_info['source_name'] = source_ref
                activity_info['source_type'] = self._map_to_database_type(source_ref)
            
            # Extract sink information  
            if 'outputs' in activity and activity['outputs']:
                sink_ref = activity['outputs'][0].get('referenceName', '')
                activity_info['sink_name'] = sink_ref
                activity_info['sink_type'] = self._map_to_database_type(sink_ref)
            
            # Extract additional type properties for more context
            type_properties = activity.get('typeProperties', {})
            
            # For Copy activities, get source and sink types
            if 'source' in type_properties:
                source_info = type_properties['source']
                source_type = source_info.get('type', '')
                if source_type and not activity_info['source_type']:
                    activity_info['source_type'] = self._convert_source_sink_type(source_type)
                
                # Extract query for Oracle sources
                if 'oracleReaderQuery' in source_info:
                    activity_info['source_query'] = source_info['oracleReaderQuery'][:100] + "..." if len(source_info['oracleReaderQuery']) > 100 else source_info['oracleReaderQuery']
            
            if 'sink' in type_properties:
                sink_info = type_properties['sink']
                sink_type = sink_info.get('type', '')
                if sink_type and not activity_info['sink_type']:
                    activity_info['sink_type'] = self._convert_source_sink_type(sink_type)
            
            # For ExecuteDataFlow activities
            if 'dataflow' in type_properties:
                dataflow_ref = type_properties['dataflow'].get('referenceName', '')
                activity_info['dataflow_name'] = dataflow_ref
            
            # For staging information
            if 'staging' in type_properties:
                staging_info = type_properties['staging']
                if 'linkedService' in staging_info:
                    staging_service = staging_info['linkedService'].get('referenceName', '')
                    activity_info['staging_service'] = staging_service
                    activity_info['staging_type'] = self._map_to_database_type(staging_service)
            
            # For compute information
            if 'compute' in type_properties:
                compute_info = type_properties['compute']
                activity_info['compute_cores'] = compute_info.get('coreCount', '')
                activity_info['compute_type'] = compute_info.get('computeType', '')
            
            extracted_data.append(activity_info)
        
        return extracted_data
        
    except Exception as e:
        return [{'error': f'Error processing pipeline: {str(e)}'}]

def _convert_source_sink_type(self, type_string):
    """Convert ADF source/sink type to readable format"""
    type_mapping = {
        'OracleSource': 'Oracle',
        'SqlServerSource': 'SQL Server',
        'ParquetSink': 'Parquet (Azure Data Lake)',
        'BlobSink': 'Azure Blob Storage',
        'AzureBlobFSSink': 'Azure Data Lake Storage Gen2',
        'DelimitedTextSink': 'Delimited Text (Azure Data Lake)',
    }
    
    return type_mapping.get(type_string, type_string)

def _map_to_database_type(self, reference_name):
    """Map reference name to database type using fuzzy matching"""
    if not reference_name:
        return 'Unknown'
    
    # Direct match
    if reference_name in self.db_mapping:
        return self.db_mapping[reference_name]
    
    # Fuzzy match
    closest_matches = get_close_matches(reference_name, self.db_mapping.keys(), n=1, cutoff=0.6)
    if closest_matches:
        return self.db_mapping[closest_matches[0]]
    
    # Pattern matching for common cases
    reference_upper = reference_name.upper()
    
    if any(pattern in reference_upper for pattern in ['ORACLE', 'ORA', 'AGLPRD', 'AGILIST']):
        return 'Oracle'
    elif any(pattern in reference_upper for pattern in ['SQL', 'MSSQL', 'SQLSERVER']):
        return 'SQL Server'
    elif any(pattern in reference_upper for pattern in ['AZURE', 'DATALAKE', 'ADLS', 'BLOB', 'SYNAPSE']):
        return 'Azure Data Lake Storage Gen2'
    elif any(pattern in reference_upper for pattern in ['PARQUET', 'CSV', 'JSON']):
        return 'Azure Data Lake Storage Gen2'
    
    return 'Unknown'

    def is_supported_file(self, file_path):
        """Check if file format is supported"""
        return Path(file_path).suffix.lower() in self.supported_extensions
    
    def extract_content(self, file_path):
        """Extract text content from various file formats"""
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.txt':
                return self._extract_from_txt(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension in ['.xlsx', '.xls']:
                return self._extract_from_excel(file_path)
            elif file_extension == '.csv':
                return self._extract_from_csv(file_path)
            elif file_extension == '.json':
                return self._extract_from_json(file_path)
            else:
                return None
                
        except Exception as e:
            print(f"Error extracting content from {file_path}: {str(e)}")
            return None

    def process_document(self, uploaded_file):
    """
    Process an uploaded document
    
    Args:
        uploaded_file: Streamlit uploaded file object
        
    Returns:
        Dict containing processing results
    """
    try:
        # Get file info
        file_details = {
            "filename": uploaded_file.name,
            "filetype": uploaded_file.type,
            "filesize": uploaded_file.size
        }
        
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            temp_path = tmp_file.name
        
        # Process based on file type
        if uploaded_file.type == "application/pdf":
            content = self._process_pdf(temp_path)
        elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            content = self._process_docx(temp_path)
        elif uploaded_file.type == "text/plain":
            content = self._process_txt(uploaded_file)
        elif uploaded_file.name.endswith('.json'):
            json_content = json.loads(uploaded_file.getvalue().decode("utf-8"))
            pipeline_data = self.extract_pipeline_data(json_content)
            return {
                "status": "success",
                "file_details": file_details,
                "pipeline_data": pipeline_data,
                "content_type": "pipeline_json"
            }
        else:
            content = "Unsupported file format"
        
        # Clean up temporary file
        os.unlink(temp_path)
        
        return {
            "status": "success",
            "file_details": file_details,
            "content_preview": content[:500] + "..." if len(content) > 500 else content,
            "content_length": len(content)
        }
        
    except Exception as e:
        return {
            "status": "error",
            "error": str(e)
        }
        
    def _detect_encoding(self, file_path):
        """Detect file encoding"""
        try:
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                result = chardet.detect(raw_data)
                return result['encoding'] or 'utf-8'
        except:
            return 'utf-8'
    
    def _extract_from_txt(self, file_path):
        """Extract content from text files"""
        try:
            encoding = self._detect_encoding(file_path)
            with open(file_path, 'r', encoding=encoding, errors='ignore') as file:
                return file.read()
        except Exception as e:
            print(f"Error reading text file: {str(e)}")
            return None
    
    def _extract_from_docx(self, file_path):
        """Extract content from Word documents"""
        try:
            doc = Document(file_path)
            content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        row_content.append(cell.text.strip())
                    table_content.append(" | ".join(row_content))
                content.extend(table_content)
            
            return "\n".join(content)
        except Exception as e:
            print(f"Error reading Word document: {str(e)}")
            return None
    
    def _extract_from_pdf(self, file_path):
        """Extract content from PDF files"""
        try:
            content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        content.append(text)
            return "\n".join(content)
        except Exception as e:
            print(f"Error reading PDF file: {str(e)}")
            return None

    def _extract_from_excel(self, file_path):
        """Extract content from Excel files"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            content = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Add sheet name
                content.append(f"Sheet: {sheet_name}")
                
                # Add column headers
                content.append(" | ".join(df.columns.astype(str)))
                
                # Add data rows (limit to first 100 rows to avoid huge content)
                for _, row in df.head(100).iterrows():
                    row_text = " | ".join(row.astype(str))
                    content.append(row_text)
                
                content.append("")  # Empty line between sheets
            
            return "\n".join(content)
        except Exception as e:
            print(f"Error reading Excel file: {str(e)}")
            return None
    
    def _extract_from_csv(self, file_path):
        """Extract content from CSV files"""
        try:
            encoding = self._detect_encoding(file_path)
            df = pd.read_csv(file_path, encoding=encoding)
            
            content = []
            # Add column headers
            content.append(" | ".join(df.columns.astype(str)))
            
            # Add data rows (limit to first 100 rows)
            for _, row in df.head(100).iterrows():
                row_text = " | ".join(row.astype(str))
                content.append(row_text)
            
            return "\n".join(content)
        except Exception as e:
            print(f"Error reading CSV file: {str(e)}")
            return None
    
    def _extract_from_json(self, file_path):
        """Extract content from JSON files, with special handling for pipeline configurations"""
        try:
            encoding = self._detect_encoding(file_path)
            with open(file_path, 'r', encoding=encoding, errors='ignore') as file:
                json_data = json.load(file)
            
            # Store the raw JSON data as a special marker for pipeline processing
            filename = Path(file_path).name
            
            # Check if this is a pipeline JSON with properties.activities structure
            if isinstance(json_data, dict) and 'properties' in json_data:
                # Return both readable text and raw JSON structure
                pipeline_name = json_data.get('name', filename.replace('.json', ''))
                
                # Create readable summary
                content_parts = [
                    f"PIPELINE_JSON_START",
                    f"JSON File: {filename}",
                    f"Azure Data Factory/Synapse Pipeline Configuration:",
                    f"Pipeline Name: {pipeline_name}",
                    json.dumps(json_data, indent=2),  # Include raw JSON
                    f"PIPELINE_JSON_END"
                ]
                
                return "\n".join(content_parts)
            
            else:
                # For non-pipeline JSON, provide a general summary
                content_parts = [f"JSON File: {filename}", "\nJSON Content Summary:"]
                if isinstance(json_data, dict):
                    for key, value in list(json_data.items())[:10]:  # Limit to first 10 keys
                        if isinstance(value, (str, int, float, bool)):
                            content_parts.append(f"{key}: {value}")
                        elif isinstance(value, list):
                            content_parts.append(f"{key}: [Array with {len(value)} items]")
                        elif isinstance(value, dict):
                            content_parts.append(f"{key}: [Object with {len(value)} properties]")
                elif isinstance(json_data, list):
                    content_parts.append(f"Array with {len(json_data)} items")
                
                return "\n".join(content_parts)
            
        except Exception as e:
            print(f"Error reading JSON file: {str(e)}")
            return f"JSON file detected but could not be parsed: {str(e)}"
    
    def extract_from_zip(self, zip_path, extract_dir=None):
        """Extract all supported files from a zip archive"""
        if extract_dir is None:
            extract_dir = tempfile.mkdtemp()
        
        extracted_files = []
        
        try:
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(extract_dir)
            
            # Process all extracted files
            for root, dirs, files in os.walk(extract_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    if self.is_supported_file(file_path):
                        content = self.extract_content(file_path)
                        if content:
                            extracted_files.append({
                                'filename': file,
                                'path': file_path,
                                'content': content
                            })
            
            return extracted_files
        except Exception as e:
            print(f"Error extracting zip file: {str(e)}")
            return []


from docx import Document

def extract_sources_from_docx(docx_path):
    """Extract source/sink names from Section 2 table in the documentation."""
    doc = Document(docx_path)
    names = []
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if cells and cells[0] and cells[0] not in ("Source/Sink Name", "Database", "Data Warehouse"):
                names.append(cells[0])
    return list(set(names))
